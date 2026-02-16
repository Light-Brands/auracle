#!/usr/bin/env python3
"""
Convert the Volume 1 complete manuscript from Markdown to a professionally
formatted DOCX file suitable for publication.

Handles: headings (H1-H4), bold, italic, bold-italic, blockquotes,
bullet lists, numbered lists, tables, horizontal rules, links, and emojis.
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ── Configuration ──────────────────────────────────────────────────────────

MANUSCRIPT_PATH = os.path.join(os.path.dirname(__file__), "vol-1-complete-manuscript.md")
OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "vol-1-see.docx")

FONT_BODY = "Georgia"
FONT_HEADING = "Georgia"
FONT_SIZE_BODY = Pt(11)
FONT_SIZE_H1 = Pt(24)
FONT_SIZE_H2 = Pt(18)
FONT_SIZE_H3 = Pt(14)
FONT_SIZE_H4 = Pt(12)
FONT_SIZE_BLOCKQUOTE = Pt(11)

LINE_SPACING = 1.15
MARGIN_TOP = Inches(0.8)
MARGIN_BOTTOM = Inches(0.8)
MARGIN_LEFT = Inches(0.9)
MARGIN_RIGHT = Inches(0.9)

COLOR_HEADING = RGBColor(0x1A, 0x1A, 0x2E)
COLOR_BODY = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_BLOCKQUOTE = RGBColor(0x44, 0x44, 0x44)
COLOR_MUTED = RGBColor(0x66, 0x66, 0x66)


# ── Title page H1s that get page breaks ────────────────────────────────────

# These H1 patterns trigger a page break before them
PAGE_BREAK_PATTERNS = [
    r"^Chapter \d+",
    r"^PART \d+",
    r"^Pause and Integrate",
    r"^Appendix [A-Z]",
    r"^About the Author",
    r"^Category [A-H]",
    r"^Card 0",
    r"^Tactical Summary",
    r"^The 48-Hour",
    r"^Quick Reference",
]


def should_page_break(text):
    """Check if this H1 heading should have a page break before it."""
    clean = text.strip()
    for pat in PAGE_BREAK_PATTERNS:
        if re.match(pat, clean):
            return True
    return False


# ── Document setup ─────────────────────────────────────────────────────────

def setup_document():
    """Create and configure the DOCX document with proper styles."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(6)
    section.page_height = Inches(9)
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT

    # Configure default style
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_BODY
    font.size = FONT_SIZE_BODY
    font.color.rgb = COLOR_BODY
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)

    # Heading 1 style
    h1 = doc.styles["Heading 1"]
    h1.font.name = FONT_HEADING
    h1.font.size = FONT_SIZE_H1
    h1.font.bold = True
    h1.font.color.rgb = COLOR_HEADING
    h1.paragraph_format.space_before = Pt(36)
    h1.paragraph_format.space_after = Pt(18)
    h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Heading 2 style
    h2 = doc.styles["Heading 2"]
    h2.font.name = FONT_HEADING
    h2.font.size = FONT_SIZE_H2
    h2.font.bold = True
    h2.font.color.rgb = COLOR_HEADING
    h2.paragraph_format.space_before = Pt(24)
    h2.paragraph_format.space_after = Pt(12)
    h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Heading 3 style
    h3 = doc.styles["Heading 3"]
    h3.font.name = FONT_HEADING
    h3.font.size = FONT_SIZE_H3
    h3.font.bold = True
    h3.font.color.rgb = COLOR_HEADING
    h3.paragraph_format.space_before = Pt(18)
    h3.paragraph_format.space_after = Pt(8)

    # Heading 4 style
    h4 = doc.styles["Heading 4"]
    h4.font.name = FONT_HEADING
    h4.font.size = FONT_SIZE_H4
    h4.font.bold = True
    h4.font.color.rgb = COLOR_HEADING
    h4.paragraph_format.space_before = Pt(12)
    h4.paragraph_format.space_after = Pt(6)

    return doc


# ── Inline formatting ──────────────────────────────────────────────────────

def add_formatted_runs(paragraph, text, base_font_size=None, base_color=None, base_italic=False):
    """Parse inline markdown (bold, italic, bold-italic, links) and add runs."""
    if base_font_size is None:
        base_font_size = FONT_SIZE_BODY
    if base_color is None:
        base_color = COLOR_BODY

    # Tokenize inline markdown
    # Order matters: bold-italic first, then bold, then italic, then links
    # Pattern: ***text*** | **text** | *text* | [text](url)
    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*)'     # bold-italic
        r'|(\*\*(.+?)\*\*)'         # bold
        r'|(\*(.+?)\*)'             # italic
        r'|(\[([^\]]+)\]\(([^)]+)\))'  # link
    )

    pos = 0
    for m in pattern.finditer(text):
        # Add plain text before match
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            run.font.name = FONT_BODY
            run.font.size = base_font_size
            run.font.color.rgb = base_color
            if base_italic:
                run.font.italic = True

        if m.group(2):  # bold-italic
            run = paragraph.add_run(m.group(2))
            run.font.name = FONT_BODY
            run.font.size = base_font_size
            run.font.bold = True
            run.font.italic = True
            run.font.color.rgb = base_color
        elif m.group(4):  # bold
            run = paragraph.add_run(m.group(4))
            run.font.name = FONT_BODY
            run.font.size = base_font_size
            run.font.bold = True
            run.font.color.rgb = base_color
            if base_italic:
                run.font.italic = True
        elif m.group(6):  # italic
            run = paragraph.add_run(m.group(6))
            run.font.name = FONT_BODY
            run.font.size = base_font_size
            run.font.italic = True
            run.font.color.rgb = base_color
        elif m.group(8):  # link
            link_text = m.group(9)
            run = paragraph.add_run(link_text)
            run.font.name = FONT_BODY
            run.font.size = base_font_size
            run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
            run.font.underline = True

        pos = m.end()

    # Add remaining text
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.name = FONT_BODY
        run.font.size = base_font_size
        run.font.color.rgb = base_color
        if base_italic:
            run.font.italic = True


# ── Title page ─────────────────────────────────────────────────────────────

def add_title_page(doc, title, subtitle, author):
    """Create a formatted title page."""
    # Spacer
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.name = FONT_HEADING
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = COLOR_HEADING
    p.paragraph_format.space_after = Pt(12)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.name = FONT_HEADING
    run.font.size = Pt(16)
    run.font.color.rgb = COLOR_MUTED
    run.font.italic = True
    p.paragraph_format.space_after = Pt(48)

    # Decorative line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("———")
    run.font.name = FONT_HEADING
    run.font.size = Pt(18)
    run.font.color.rgb = COLOR_MUTED
    p.paragraph_format.space_after = Pt(48)

    # Author
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(author)
    run.font.name = FONT_HEADING
    run.font.size = Pt(18)
    run.font.color.rgb = COLOR_HEADING
    p.paragraph_format.space_after = Pt(12)

    # Series line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("The Sovereignty Series — Volume 1")
    run.font.name = FONT_HEADING
    run.font.size = Pt(12)
    run.font.color.rgb = COLOR_MUTED
    run.font.italic = True

    # Page break after title
    doc.add_page_break()


# ── Table handling ─────────────────────────────────────────────────────────

def add_table(doc, rows):
    """Add a formatted table to the document."""
    if len(rows) < 2:
        return

    # Parse header and data
    header = [cell.strip() for cell in rows[0].strip("|").split("|")]
    # Skip separator row (row 1)
    data_rows = []
    for row in rows[2:]:
        cells = [cell.strip() for cell in row.strip("|").split("|")]
        data_rows.append(cells)

    num_cols = len(header)
    table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Style header row
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = FONT_BODY
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Dark background for header
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1A1A2E"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row_data in enumerate(data_rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= num_cols:
                break
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            add_formatted_runs(p, cell_text, base_font_size=Pt(10))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Alternating row colors
            if r_idx % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    # Add spacing after table
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


# ── Horizontal rule ────────────────────────────────────────────────────────

def add_horizontal_rule(doc):
    """Add a decorative separator."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("•  •  •")
    run.font.name = FONT_HEADING
    run.font.size = Pt(12)
    run.font.color.rgb = COLOR_MUTED
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)


# ── Page break ─────────────────────────────────────────────────────────────

def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()


# ── Main conversion ───────────────────────────────────────────────────────

def convert_manuscript(input_path, output_path):
    """Convert the full markdown manuscript to DOCX."""
    print(f"Reading manuscript: {input_path}")
    with open(input_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    print(f"  {len(lines)} lines read")

    doc = setup_document()

    # Add title page
    add_title_page(
        doc,
        "You Are Not Crazy",
        "The Narcissism Decoder:\n60 Manipulation Patterns They Used Against You",
        "Jennifer Brooke Lawless"
    )

    i = 0
    is_first_h1 = True  # Skip first H1 (title, already on title page)
    table_buffer = []
    in_table = False
    prev_was_blank = False

    while i < len(lines):
        line = lines[i].rstrip("\n")

        # ── Table detection ────────────────────────────────────────────
        if line.startswith("|") and "|" in line[1:]:
            table_buffer.append(line)
            in_table = True
            i += 1
            continue
        elif in_table:
            # End of table
            add_table(doc, table_buffer)
            table_buffer = []
            in_table = False
            # Don't increment i; process current line normally

        # ── Blank lines ────────────────────────────────────────────────
        if line.strip() == "":
            prev_was_blank = True
            i += 1
            continue

        # ── Horizontal rules ──────────────────────────────────────────
        if re.match(r"^---+$", line.strip()) or re.match(r"^\*\*\*+$", line.strip()):
            add_horizontal_rule(doc)
            prev_was_blank = False
            i += 1
            continue

        # ── H1 ─────────────────────────────────────────────────────────
        m = re.match(r"^# (.+)$", line)
        if m:
            heading_text = m.group(1).strip()
            if is_first_h1:
                # Skip the very first H1 (book title) — already on title page
                is_first_h1 = False
                i += 1
                continue

            # Page break before major sections
            if should_page_break(heading_text):
                add_page_break(doc)

            p = doc.add_heading(heading_text, level=1)
            # Re-apply font to heading runs (python-docx quirk)
            for run in p.runs:
                run.font.name = FONT_HEADING
                run.font.size = FONT_SIZE_H1
                run.font.color.rgb = COLOR_HEADING
            prev_was_blank = False
            i += 1
            continue

        # ── H2 ─────────────────────────────────────────────────────────
        m = re.match(r"^## (.+)$", line)
        if m:
            heading_text = m.group(1).strip()

            # Skip the first H2 if it's the subtitle (already on title page)
            if heading_text.startswith("The Narcissism Decoder"):
                i += 1
                continue

            p = doc.add_heading(heading_text, level=2)
            for run in p.runs:
                run.font.name = FONT_HEADING
                run.font.size = FONT_SIZE_H2
                run.font.color.rgb = COLOR_HEADING
            prev_was_blank = False
            i += 1
            continue

        # ── H3 ─────────────────────────────────────────────────────────
        m = re.match(r"^### (.+)$", line)
        if m:
            heading_text = m.group(1).strip()
            p = doc.add_heading(heading_text, level=3)
            for run in p.runs:
                run.font.name = FONT_HEADING
                run.font.size = FONT_SIZE_H3
                run.font.color.rgb = COLOR_HEADING
            prev_was_blank = False
            i += 1
            continue

        # ── H4 ─────────────────────────────────────────────────────────
        m = re.match(r"^#### (.+)$", line)
        if m:
            heading_text = m.group(1).strip()
            p = doc.add_heading(heading_text, level=4)
            for run in p.runs:
                run.font.name = FONT_HEADING
                run.font.size = FONT_SIZE_H4
                run.font.color.rgb = COLOR_HEADING
            prev_was_blank = False
            i += 1
            continue

        # ── Blockquote ─────────────────────────────────────────────────
        if line.startswith("> "):
            quote_text = line[2:].strip()
            # Collect multi-line blockquotes
            while i + 1 < len(lines) and lines[i + 1].startswith("> "):
                i += 1
                quote_text += " " + lines[i].rstrip("\n")[2:].strip()

            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.right_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            # Add left border via XML
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:left w:val="single" w:sz="12" w:space="8" w:color="999999"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            add_formatted_runs(p, quote_text, base_font_size=FONT_SIZE_BLOCKQUOTE,
                               base_color=COLOR_BLOCKQUOTE, base_italic=True)
            prev_was_blank = False
            i += 1
            continue

        # ── Numbered list ──────────────────────────────────────────────
        m = re.match(r"^(\d+)\.\s+(.+)$", line)
        if m:
            list_text = m.group(2).strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.space_before = Pt(1)
            run = p.add_run(f"{m.group(1)}. ")
            run.font.name = FONT_BODY
            run.font.size = FONT_SIZE_BODY
            run.font.bold = True
            run.font.color.rgb = COLOR_BODY
            add_formatted_runs(p, list_text)
            prev_was_blank = False
            i += 1
            continue

        # ── Bullet list ────────────────────────────────────────────────
        m = re.match(r"^[-*]\s+(.+)$", line)
        if m:
            list_text = m.group(1).strip()
            # Check for sub-items on next lines
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.space_before = Pt(1)
            run = p.add_run("•  ")
            run.font.name = FONT_BODY
            run.font.size = FONT_SIZE_BODY
            run.font.color.rgb = COLOR_MUTED
            add_formatted_runs(p, list_text)
            prev_was_blank = False
            i += 1
            continue

        # ── Sub-bullet (indented) ──────────────────────────────────────
        m = re.match(r"^\s{2,}[-*]\s+(.+)$", line)
        if m:
            list_text = m.group(1).strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.7)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(1)
            run = p.add_run("◦  ")
            run.font.name = FONT_BODY
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR_MUTED
            add_formatted_runs(p, list_text, base_font_size=Pt(10.5))
            prev_was_blank = False
            i += 1
            continue

        # ── Regular paragraph ──────────────────────────────────────────
        # Collect continuation lines (non-blank, non-special)
        para_text = line.strip()

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        add_formatted_runs(p, para_text)
        prev_was_blank = False
        i += 1

    # Flush any remaining table
    if table_buffer:
        add_table(doc, table_buffer)

    # Save
    print(f"Saving DOCX: {output_path}")
    doc.save(output_path)
    print(f"Done! File saved to: {output_path}")

    # Stats
    file_size = os.path.getsize(output_path)
    print(f"  File size: {file_size / 1024 / 1024:.1f} MB")


if __name__ == "__main__":
    convert_manuscript(MANUSCRIPT_PATH, OUTPUT_PATH)
