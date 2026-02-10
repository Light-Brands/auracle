#!/usr/bin/env python3
"""Convert vol 1 complete manuscript markdown to Word .docx format."""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT


def parse_markdown_to_docx(md_path, docx_path):
    """Parse markdown file and create a formatted Word document."""
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Georgia"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Configure heading styles
    for level in range(1, 5):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.name = "Georgia"
        heading_style.font.color.rgb = RGBColor(0x2C, 0x2C, 0x2C)
        if level == 1:
            heading_style.font.size = Pt(24)
            heading_style.paragraph_format.space_before = Pt(36)
            heading_style.paragraph_format.space_after = Pt(18)
            heading_style.paragraph_format.page_break_before = True
        elif level == 2:
            heading_style.font.size = Pt(18)
            heading_style.paragraph_format.space_before = Pt(24)
            heading_style.paragraph_format.space_after = Pt(12)
        elif level == 3:
            heading_style.font.size = Pt(14)
            heading_style.paragraph_format.space_before = Pt(18)
            heading_style.paragraph_format.space_after = Pt(8)
        elif level == 4:
            heading_style.font.size = Pt(12)
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)

    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    lines = content.split("\n")
    i = 0
    in_blockquote = False
    blockquote_text = []
    in_list = False
    first_h1 = True

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip HTML div tags (alignment markers from markdown)
        if stripped.startswith("<div") or stripped.startswith("</div>"):
            i += 1
            continue

        # Skip empty lines
        if not stripped:
            if in_blockquote and blockquote_text:
                # End blockquote
                add_blockquote(doc, "\n".join(blockquote_text))
                blockquote_text = []
                in_blockquote = False
            in_list = False
            i += 1
            continue

        # Horizontal rules -> thin line or page break context
        if stripped in ("---", "***", "___"):
            if in_blockquote and blockquote_text:
                add_blockquote(doc, "\n".join(blockquote_text))
                blockquote_text = []
                in_blockquote = False
            # Add a subtle separator
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("─" * 30)
            run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
            run.font.size = Pt(8)
            i += 1
            continue

        # Headings
        heading_match = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if heading_match:
            if in_blockquote and blockquote_text:
                add_blockquote(doc, "\n".join(blockquote_text))
                blockquote_text = []
                in_blockquote = False
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            text = clean_markdown_inline(text)

            # First H1 is the title - special treatment (no page break)
            if level == 1 and first_h1:
                first_h1 = False
                p = doc.add_heading(text, level=1)
                p.paragraph_format.page_break_before = False
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p = doc.add_heading(text, level=level)
            i += 1
            continue

        # Blockquotes
        if stripped.startswith(">"):
            in_blockquote = True
            quote_text = stripped.lstrip("> ").strip()
            if quote_text:
                blockquote_text.append(quote_text)
            i += 1
            continue

        # Unordered list items
        list_match = re.match(r"^(\s*)[-*+]\s+(.*)", line)
        if list_match:
            if in_blockquote and blockquote_text:
                add_blockquote(doc, "\n".join(blockquote_text))
                blockquote_text = []
                in_blockquote = False
            indent_level = len(list_match.group(1)) // 2
            text = list_match.group(2).strip()
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_text(p, text)
            p.paragraph_format.left_indent = Inches(0.25 + (indent_level * 0.25))
            in_list = True
            i += 1
            continue

        # Ordered list items
        olist_match = re.match(r"^(\s*)\d+\.\s+(.*)", line)
        if olist_match:
            if in_blockquote and blockquote_text:
                add_blockquote(doc, "\n".join(blockquote_text))
                blockquote_text = []
                in_blockquote = False
            indent_level = len(olist_match.group(1)) // 2
            text = olist_match.group(2).strip()
            p = doc.add_paragraph(style="List Number")
            add_formatted_text(p, text)
            p.paragraph_format.left_indent = Inches(0.25 + (indent_level * 0.25))
            in_list = True
            i += 1
            continue

        # Table rows (simple handling)
        if stripped.startswith("|"):
            if in_blockquote and blockquote_text:
                add_blockquote(doc, "\n".join(blockquote_text))
                blockquote_text = []
                in_blockquote = False
            # Collect all table rows
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = lines[i].strip()
                # Skip separator rows
                if not re.match(r"^\|[\s\-:|]+\|$", row):
                    cells = [c.strip() for c in row.split("|")[1:-1]]
                    table_rows.append(cells)
                i += 1
            if table_rows:
                add_table(doc, table_rows)
            continue

        # Regular paragraph
        if in_blockquote and blockquote_text:
            add_blockquote(doc, "\n".join(blockquote_text))
            blockquote_text = []
            in_blockquote = False

        p = doc.add_paragraph()
        add_formatted_text(p, stripped)
        i += 1

    # Flush any remaining blockquote
    if blockquote_text:
        add_blockquote(doc, "\n".join(blockquote_text))

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


def add_blockquote(doc, text):
    """Add a formatted blockquote to the document."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    add_formatted_text(p, text, italic_default=True, color=RGBColor(0x55, 0x55, 0x55))


def add_table(doc, rows):
    """Add a simple table to the document."""
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Light Grid Accent 1"
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            if ci < num_cols:
                cell = table.cell(ri, ci)
                cell.text = clean_markdown_inline(cell_text)
                for paragraph in cell.paragraphs:
                    paragraph.style = doc.styles["Normal"]
                    for run in paragraph.runs:
                        run.font.size = Pt(10)


def add_formatted_text(paragraph, text, italic_default=False, color=None):
    """Add text with inline markdown formatting (bold, italic, bold+italic)."""
    # Process inline formatting: ***bold italic***, **bold**, *italic*, `code`
    pattern = re.compile(
        r"(\*\*\*(.+?)\*\*\*"  # ***bold italic***
        r"|\*\*(.+?)\*\*"       # **bold**
        r"|\*(.+?)\*"           # *italic*
        r"|_(.+?)_"             # _italic_
        r"|`(.+?)`"             # `code`
        r"|([^*_`]+))"          # plain text
    )

    for match in pattern.finditer(text):
        if match.group(2):  # bold italic
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(3):  # bold
            run = paragraph.add_run(match.group(3))
            run.bold = True
            if italic_default:
                run.italic = True
        elif match.group(4):  # italic with *
            run = paragraph.add_run(match.group(4))
            run.italic = True
        elif match.group(5):  # italic with _
            run = paragraph.add_run(match.group(5))
            run.italic = True
        elif match.group(6):  # code
            run = paragraph.add_run(match.group(6))
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x66, 0x33, 0x99)
        elif match.group(7):  # plain text
            run = paragraph.add_run(match.group(7))
            if italic_default:
                run.italic = True

        if color and run:
            run.font.color.rgb = color


def clean_markdown_inline(text):
    """Remove markdown inline formatting for plain text contexts."""
    text = re.sub(r"\*\*\*(.+?)\*\*\*", r"\1", text)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"_(.+?)_", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


if __name__ == "__main__":
    md_path = "/home/user/auracle/book/sovereignty-series/vol-1-decoder/vol-1-complete-manuscript.md"
    docx_path = "/home/user/auracle/book/sovereignty-series/vol-1-decoder/vol-1-complete-manuscript.docx"
    parse_markdown_to_docx(md_path, docx_path)
